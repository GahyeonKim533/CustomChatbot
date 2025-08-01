import os
import streamlit as st
from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain.memory import StreamlitChatMessageHistory
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from dotenv import load_dotenv
from docx import Document
from io import BytesIO

# 환경 변수 로드
load_dotenv()
API_KEY = os.getenv("OPENAI_API_KEY")
os.environ["OPENAI_API_KEY"] = API_KEY

# Streamlit 설정
st.set_page_config(layout="wide", page_title="자기소개서/이력서 작성 앱")
st.title('자기소개서/이력서 작성 앱')

# Streamlit 기반 메시지 히스토리 (세션 저장형)
history = StreamlitChatMessageHistory()

# LangChain 설정
llm = ChatOpenAI(model='gpt-4o-mini', temperature=0.1)

# 새로운 답변 생성 프롬프트
qa_prompt = ChatPromptTemplate.from_messages(
    [
        ("system", """
            당신은 "인터뷰 & 이력서 & 자기소개서 코치 GPT"입니다. 당신의 주요 역할은 다음과 같습니다:

1. 사용자의 이력서/자기소개서를 개선하거나 새로 작성하도록 돕습니다.
2. 사용자가 지원하는 직무에 맞춰 모의 면접을 진행하며, 질문-답변-피드백-다음 질문의 순서를 따릅니다.
3. IT 직군, 대학원/박사 과정 진학, 신입 및 경력직 취업 준비자를 주 대상으로 합니다.

작업 방식은 다음과 같습니다:

- 사용자가 이력서/자소서 최적화를 요청한 경우:
  1. 먼저 5개 이상의 구체적인 질문을 통해 사용자 정보를 수집합니다.
  2. 수집된 정보를 바탕으로 구체적인 수정 제안 또는 개선된 텍스트를 작성합니다.
  3. 필요시 DOCX 문서 형식으로 생성하여 제공합니다.
  4. 항상 문서나 내용 생성 전에 사용자에게 한 질문에 하나씩 물어보고 대답을 기다립니다.

- 사용자가 이력서를 새로 만들고 싶다고 한 경우:
  1. 최소 6개의 질문을 단계별로 하며 사용자 정보를 수집합니다.
  2. 우선 사용자의 이름 및 주소, 연락처를 수집하고, 자소서를 이미 만들었다면 자소서를 참조합니다.
  3. 이를 바탕으로 맞춤형 이력서를 처음부터 생성합니다.
  4. 이력서를 생성했을때는 이력서 내용만을 출력합니다.
  5. 필요시 DOCX 문서로 제공합니다.

- 사용자가 자기소개서를 새로 만들고 싶다고 한 경우:
  1. 최소 5개의 질문을 단계별로 하며 사용자 정보를 수집합니다.
  2. 이력서를 이미 만들었다면 이력서를 참조합니다.
  3. 도전정신/책임의식 등 기업이 추구하는 인재상 위주로 키워드를 질문합니다.
  4. 이를 바탕으로 맞춤형 자기소개서를 1000자 내외로 생성합니다.
  5. 자기소개서를 생성했을때는 자기소개서 내용만을 출력합니다.
  6. 필요시 DOCX 문서로 제공합니다.

- 사용자가 면접 연습을 요청한 경우:
  1. 먼저 자기소개 또는 이력서를 요청합니다. 앞서 새로 만든 자소서나 이력서가 있다면 이를 참조합니다. 없다고 할 경우 넘어갑니다.
  2. 그 다음, 사용자가 지원하는 직무명과 JD(직무 요건)를 요청합니다.
  3. 이후, 당신은 인터뷰어 역할로 전환되어 질문을 하나씩 제시하고 답변을 기다립니다.
  4. 사용자의 답변을 분석하여 장점, 개선점, 모범답안을 제시합니다.
  5. 다음 질문으로 넘어가며 반복합니다.
  6. 면접 질문은 해당 직무의 특성(예: 기술직, 제품직, 경영직 등)에 따라 맞춤화합니다.

추가 규칙:
- 한 번에 하나의 질문만 하세요.
- 설명은 친절하고 따뜻하게, 하지만 구체적이고 실용적으로 작성하세요.
- 사용자가 사용한 언어(예: 영어, 한국어 등)를 그대로 따라 사용하세요.
"""
),
        MessagesPlaceholder(variable_name="chat_history"),
        ("human", "{input}"),
    ]
)

# LCEL 형식으로 체인 구성
conversation_chain = (
    RunnablePassthrough.assign(
        chat_history=lambda x: history.messages
    )
    | qa_prompt
    | llm
    | StrOutputParser()
)

# GPT 응답 함수
def ask_gpt(user_input: str) -> str:
    response = conversation_chain.invoke({
        "input": user_input,
    })
    history.add_user_message(user_input)
    history.add_ai_message(response)
    return response

# DOCX 파일 생성 함수
def create_docx_file(content: str, filename: str = "문서.docx"):
    doc = Document()
    doc.add_heading(filename.replace(".docx", ""), level=1)
    doc.add_paragraph(content)
    
    # BytesIO를 사용해 메모리 상에서 파일 생성
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# 이전 메시지 출력
for msg in history.messages:
    role = "user" if msg.type == "human" else "assistant"
    with st.chat_message(role):
        st.markdown(msg.content)

# 사용자 입력받기
if user_question := st.chat_input("질문을 입력하세요"):
    with st.chat_message("user"):
        st.markdown(user_question)
    with st.spinner("AI가 답변을 생성 중입니다..."):
        ai_answer = ask_gpt(user_question)
    with st.chat_message("assistant"):
        st.markdown(ai_answer)
    
    # 다운로드 버튼을 조건부로 표시
    if "DOCX" in ai_answer or "docx" in ai_answer:
        # GPT의 최종 답변을 docx 파일로 변환
        doc_file = create_docx_file(ai_answer, "작성_완료.docx")
        
        st.download_button(
            label="DOCX 파일 다운로드",
            data=doc_file,
            file_name="작성_완료.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )